from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import date

doc = Document()

# ── Page setup: Portrait A4 ──
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)

# ── Styles ──
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
MED_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
RED       = RGBColor(0xC0, 0x39, 0x2B)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY      = RGBColor(0x55, 0x55, 0x55)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.space_before = Pt(1)

for lvl, sz, clr in [(1, 16, DARK_BLUE), (2, 13, MED_BLUE), (3, 11, DARK_BLUE)]:
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = 'Calibri'
    hs.font.size = Pt(sz)
    hs.font.color.rgb = clr
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(10 if lvl == 1 else 6)
    hs.paragraph_format.space_after  = Pt(4)


def add_para(text, bold=False, italic=False, size=10, align=None, color=None, space_after=3):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * level)
    return p


def shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell(cell, text, bold=False, size=10, align=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.bold = bold
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)

def make_table(headers, data, col_widths=None):
    ncols = len(headers)
    nrows = len(data) + 1
    t = doc.add_table(rows=nrows, cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell(c, h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(c, "2E75B6")
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            c = t.rows[ri+1].cells[ci]
            set_cell(c, str(val), size=9)
            if ri % 2 == 1:
                shade_cell(c, "F2F2F2")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# ═══════════════════════════════════════════════════════════
# DOCUMENT CONTROL HEADER (table at top of every page)
# ═══════════════════════════════════════════════════════════
def add_doc_header():
    t = doc.add_table(rows=2, cols=4)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: Document ID
    set_cell(t.rows[0].cells[0], "QP-PUR-001", bold=True, size=9, color=WHITE)
    shade_cell(t.rows[0].cells[0], "1F3864")
    t.rows[0].cells[0].width = Cm(3.5)

    set_cell(t.rows[0].cells[1], "Purchasing Procedure", bold=True, size=9, color=WHITE)
    shade_cell(t.rows[0].cells[1], "1F3864")
    t.rows[0].cells[1].width = Cm(6.0)

    set_cell(t.rows[0].cells[2], "Rev. 01", size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(t.rows[0].cells[2], "1F3864")
    t.rows[0].cells[2].width = Cm(2.0)

    set_cell(t.rows[0].cells[3], f"Page 1 of ?", size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)
    shade_cell(t.rows[0].cells[3], "1F3864")
    t.rows[0].cells[3].width = Cm(3.0)

    # Row 1: Effective date
    set_cell(t.rows[1].cells[0], "Effective Date:", bold=True, size=8)
    shade_cell(t.rows[1].cells[0], "D6E4F0")
    set_cell(t.rows[1].cells[1], date.today().strftime("%d/%m/%Y"), size=8)
    shade_cell(t.rows[1].cells[1], "D6E4F0")
    set_cell(t.rows[1].cells[2], "Status:", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(t.rows[1].cells[2], "FFF2CC")
    set_cell(t.rows[1].cells[3], "[ DRAFT ]", bold=True, size=8, color=RED, align=WD_ALIGN_PARAGRAPH.RIGHT)
    shade_cell(t.rows[1].cells[3], "FFF2CC")

    doc.add_paragraph()


add_doc_header()

# ═══════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════
add_para("QUALITY PROCEDURE", bold=True, size=14, color=DARK_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Purchasing Procedure", bold=True, size=16, color=MED_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("ขั้นตอนการจัดซื้อ", bold=True, size=12, color=RGBColor(0x88, 0x88, 0x88),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(f"Document No.: QP-PUR-001     |     Revision: 01     |     "
         f"Effective Date: {date.today().strftime('%d/%m/%Y')}     |     Status: DRAFT",
         size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════
# 1. PURPOSE
# ═══════════════════════════════════════════════════════════
doc.add_heading("1. Purpose (วัตถุประสงค์)", level=1)
add_para("The purpose of this procedure is to define the process for purchasing products and services "
         "that affect product quality, ensuring they conform to specified requirements in accordance "
         "with ISO 9001:2015 clause 8.4.")

add_bullet("Establish a standardized purchasing process across the organization")
add_bullet("Ensure purchased products and services meet specified requirements")
add_bullet("Define criteria for supplier selection, evaluation, and re-evaluation")
add_bullet("Ensure traceability and transparency of all purchasing activities")

# ═══════════════════════════════════════════════════════════
# 2. SCOPE
# ═══════════════════════════════════════════════════════════
doc.add_heading("2. Scope (ขอบเขต)", level=1)
add_para("This procedure applies to all purchasing activities for materials, components, equipment, "
         "and services that affect product quality, including outsourced processes.")

add_bullet("Raw materials and production consumables")
add_bullet("Packaging materials that contact the product")
add_bullet("Equipment and spare parts affecting product quality")
add_bullet("Outsourced services (calibration, testing, maintenance, etc.)")
add_bullet("Subcontracted manufacturing processes")

add_para("Excluded: General office supplies and services that do not impact product quality "
         "(these may follow simplified procurement guidelines).", italic=True, size=9)

# ═══════════════════════════════════════════════════════════
# 3. REFERENCES
# ═══════════════════════════════════════════════════════════
doc.add_heading("3. References (เอกสารอ้างอิง)", level=1)
make_table(
    ["Doc. No.", "Title"],
    [
        ["QM-001", "Quality Manual"],
        ["WI-PUR-001", "Supplier Evaluation Work Instruction"],
        ["WI-QC-001", "Incoming Inspection Work Instruction"],
        ["FR-PR-001", "Purchase Requisition Form"],
        ["FR-PUR-001", "Purchase Order Form"],
        ["FR-PUR-002", "Supplier Evaluation Form"],
        ["FR-PUR-003", "Approved Supplier List (ASL)"],
        ["FR-QC-001", "Inspection Report Form"],
        ["FR-QC-002", "Nonconformance Report (NCR)"],
        ["FR-PUR-004", "Supplier Scorecard"],
        ["FR-PUR-005", "Supplier Corrective Action Request (SCAR)"],
        ["ISO 9001:2015", "Quality Management Systems — Requirements"],
    ],
    col_widths=[3.5, 12.0]
)

# ═══════════════════════════════════════════════════════════
# 4. DEFINITIONS
# ═══════════════════════════════════════════════════════════
doc.add_heading("4. Definitions (คำจำกัดความ)", level=1)
make_table(
    ["Term", "Definition"],
    [
        ["Supplier / ผู้ขาย", "An organization or person that provides a product or service to the company."],
        ["Approved Supplier List (ASL)", "List of suppliers who have been evaluated and approved for purchasing."],
        ["Purchase Order (PO)", "A formal document issued to a supplier specifying products/services, quantity, price, and delivery terms."],
        ["Purchase Requisition (PR)", "An internal document requesting the purchase of goods or services."],
        ["Nonconformance Report (NCR)", "A report documenting a product/service that does not meet specified requirements."],
        ["Corrective Action Request (SCAR)", "A request sent to a supplier to investigate and correct a nonconformance."],
    ],
    col_widths=[4.5, 11.0]
)

# ═══════════════════════════════════════════════════════════
# 5. RESPONSIBILITIES
# ═══════════════════════════════════════════════════════════
doc.add_heading("5. Responsibilities (ความรับผิดชอบ)", level=1)
make_table(
    ["Role", "Responsibility"],
    [
        ["Top Management", "Approve purchasing policy, approve new suppliers, approve high-value POs."],
        ["Purchasing Dept.", "Execute purchasing process, supplier selection, PO issuance, order follow-up."],
        ["Quality Control (QC)", "Inspect incoming goods, approve/reject materials, maintain quality records."],
        ["Warehouse", "Receive goods, verify quantity, store materials, notify QC for inspection."],
        ["Engineering / Production", "Define specifications, standards, and technical requirements."],
        ["Requestor", "Identify needs, prepare Purchase Requisition (PR)."],
    ],
    col_widths=[4.5, 11.0]
)

# ═══════════════════════════════════════════════════════════
# 6. PROCEDURE
# ═══════════════════════════════════════════════════════════
doc.add_heading("6. Procedure (ขั้นตอนการดำเนินงาน)", level=1)

# 6.1 Process Flow
doc.add_heading("6.1 Purchasing Process Flow", level=2)
add_para("The purchasing process consists of 13 sequential steps as follows:", size=10)

make_table(
    ["Step", "Activity", "Responsible", "Document", "ISO Clause"],
    [
        ["1", "Identify Requirement", "Requestor", "PR (FR-PR-001)", "8.4.2"],
        ["2", "Obtain Purchase Approval", "Dept. Head", "Approved PR", "5.1 / 7.1"],
        ["3", "Select Supplier", "Purchasing", "ASL (FR-PUR-003)", "8.4.1"],
        ["4", "Request Quotation (RFQ)", "Purchasing", "Quotation / RFQ", "8.4.2"],
        ["5", "Create Purchase Order", "Purchasing", "PO (FR-PUR-001)", "8.4.2"],
        ["6", "Approve PO", "Authorized Approver", "Approved PO", "5.1"],
        ["7", "Send PO to Supplier", "Purchasing", "PO Record", "8.4.2"],
        ["8", "Expedite / Follow-up", "Purchasing", "Expediting Log", "8.4.2"],
        ["9", "Receive Goods", "Warehouse", "DN / GRN", "8.4.3"],
        ["10", "Inspect Quality", "QC", "Insp. Report (FR-QC-001)", "8.4.3"],
        ["11", "Accept / Reject", "Warehouse / QC", "GRN / NCR", "8.4.3 / 8.7"],
        ["12", "Process Invoice / Pay", "Accounting", "Invoice", "—"],
        ["13", "Record & Measure", "Purchasing / QC", "Scorecard", "8.4.1 / 9.1"],
    ],
    col_widths=[1.2, 3.5, 2.5, 4.0, 2.0]
)

# 6.2 Detail
doc.add_heading("6.2 Step Details", level=2)

steps_detail = [
    ("6.2.1 Identify Requirement",
     "The Requestor identifies the need for materials or services and completes the "
     "Purchase Requisition (FR-PR-001), specifying item description, quantity, required date, "
     "specification/drawing number, and any special requirements."),
    ("6.2.2 Obtain Purchase Approval",
     "The Department Head reviews the PR for necessity, budget availability, and specification "
     "accuracy. Approved PRs are forwarded to the Purchasing Department."),
    ("6.2.3 Select Supplier",
     "Purchasing selects a supplier from the Approved Supplier List (ASL). If a new supplier is required, "
     "the evaluation process per Section 7 must be completed before purchasing."),
    ("6.2.4 Request Quotation",
     "For purchases above the threshold (e.g., 50,000 THB), at least three quotations shall be obtained. "
     "Quotations are compared for price, quality, delivery, and terms."),
    ("6.2.5 Create Purchase Order",
     "Purchasing creates a PO (FR-PUR-001) including: item description, quantity, unit price, "
     "total amount, delivery date, Incoterms, payment terms, and ISO requirements per 8.4.2."),
    ("6.2.6 Approve PO",
     "The PO must be reviewed and signed by an authorized approver based on the approval authority matrix."),
    ("6.2.7 Send PO to Supplier",
     "The approved PO is sent to the supplier via email, fax, or EDI. A confirmation copy is requested."),
    ("6.2.8 Expedite / Follow-up",
     "For critical items, Purchasing conducts regular follow-up to ensure on-time delivery. "
     "Any delays are communicated to the Requestor and recorded."),
    ("6.2.9 Receive Goods",
     "Warehouse receives the goods, verifies the Delivery Note against the PO, confirms quantity, "
     "and records receipt in GRN (Goods Received Note)."),
    ("6.2.10 Inspect Quality",
     "QC inspects the goods per the Sampling Inspection Plan. Results are recorded in FR-QC-001."),
    ("6.2.11 Accept / Reject",
     "Accepted goods are moved to inventory. Rejected goods are quarantined, and an NCR (FR-QC-002) is issued. "
     "For rejected goods, a SCAR (FR-PUR-005) may be sent to the supplier."),
    ("6.2.12 Process Payment",
     "Accounting processes the supplier invoice against the PO and GRN, then makes payment per agreed terms."),
    ("6.2.13 Record & Measure",
     "Purchasing maintains records and measures supplier performance using the Supplier Scorecard (FR-PUR-004)."),
]

for title, detail in steps_detail:
    doc.add_heading(title, level=3)
    add_para(detail)

# ═══════════════════════════════════════════════════════════
# 7. SUPPLIER EVALUATION
# ═══════════════════════════════════════════════════════════
doc.add_heading("7. Supplier Evaluation (การประเมินผู้ขาย)", level=1)

doc.add_heading("7.1 New Supplier Evaluation Criteria", level=2)
make_table(
    ["Criteria", "Description", "Weight"],
    [
        ["Quality", "Quality system (ISO 9001), process control, quality history", "35%"],
        ["Price", "Price competitiveness, cost structure transparency", "20%"],
        ["Delivery", "On-time delivery capability, flexibility, lead time", "20%"],
        ["Service", "Responsiveness, technical support, after-sales service", "10%"],
        ["Reliability", "Company history, financial stability, customer references", "10%"],
        ["Environment / Safety", "Regulatory compliance, ISO 14001 / ISO 45001", "5%"],
    ],
    col_widths=[4.0, 9.5, 2.0]
)

doc.add_heading("7.2 Evaluation Rating", level=2)
make_table(
    ["Score", "Rating", "Action"],
    [
        ["\u2265 85%", "A — Excellent", "Continuous improvement, consider preferential treatment"],
        ["70\u201384%", "B — Good", "Periodic monitoring, advise improvement areas"],
        ["50\u201369%", "C — Fair", "Require Corrective Action Plan (CAP), close monitoring"],
        ["< 50%", "D — Poor", "Consider removal from ASL / find alternative supplier"],
    ],
    col_widths=[2.5, 4.5, 8.5]
)

doc.add_heading("7.3 Re-evaluation Frequency", level=2)
make_table(
    ["Rating", "Frequency", "Method"],
    [
        ["A", "Every 12 months", "Review Scorecard + inquiry"],
        ["B", "Every 6 months", "Review Scorecard + send evaluation form"],
        ["C", "Every 3 months", "Review Scorecard + possible on-site audit"],
        ["D", "Immediately", "Execute CAP or remove from ASL"],
    ],
    col_widths=[2.0, 3.5, 10.0]
)

# ═══════════════════════════════════════════════════════════
# 8. PURCHASE ORDER REQUIREMENTS
# ═══════════════════════════════════════════════════════════
doc.add_heading("8. Purchase Order Requirements (ข้อกำหนดใบสั่งซื้อ)", level=1)
add_para("All Purchase Orders must include the following information as required by ISO 9001:2015 clause 8.4.2:")

make_table(
    ["Clause", "Requirement"],
    [
        ["8.4.2 a)", "Specifications for the product or service to be purchased (spec, drawing, standard, etc.)"],
        ["8.4.2 b)", "Approval requirements: product, procedure, process, and equipment"],
        ["8.4.2 c)", "Qualification requirements for personnel (if applicable)"],
        ["8.4.2 d)", "Quality management system requirements (e.g., ISO 9001 certification)"],
    ],
    col_widths=[2.5, 13.0]
)

# ═══════════════════════════════════════════════════════════
# 9. INCOMING INSPECTION
# ═══════════════════════════════════════════════════════════
doc.add_heading("9. Incoming Inspection (การตรวจสอบขาเข้า)", level=1)
make_table(
    ["Step", "Activity", "Responsible", "ISO Clause"],
    [
        ["1", "Receive Delivery Note & check against PO", "Warehouse", "8.4.3"],
        ["2", "Count quantity and verify documents", "Warehouse", "8.4.3"],
        ["3", "Visual inspection for damage", "Warehouse", "8.4.3"],
        ["4", "Notify QC for quality inspection", "Warehouse", "8.4.3"],
        ["5", "Perform inspection per Inspection Plan", "QC", "8.4.3"],
        ["6", "Decision: Accept / Reject / Conditional Accept", "QC", "8.4.3"],
        ["7", "Accept \u2192 stock / Reject \u2192 NCR + SCAR", "Warehouse / QC", "8.4.3 / 8.7"],
    ],
    col_widths=[1.5, 6.0, 3.0, 2.5]
)

doc.add_heading("9.1 Sampling Inspection Plan", level=2)
make_table(
    ["Material Type", "Inspection Level", "Sample Size", "AQL", "Method"],
    [
        ["Critical (A)", "100%", "All", "Zero (C=0)", "Per Spec Sheet"],
        ["Major (B)", "Random", "\u221AN or MIL-STD-1916", "AQL 1.0%", "Per Spec Sheet"],
        ["Minor (C)", "Random", "MIL-STD-1916 Level II", "AQL 4.0%", "Visual + Dim."],
        ["Services", "—", "Cert / Report review", "—", "Document review"],
    ],
    col_widths=[3.5, 2.5, 4.0, 2.5, 3.0]
)

# ═══════════════════════════════════════════════════════════
# 10. NONCONFORMING PRODUCT
# ═══════════════════════════════════════════════════════════
doc.add_heading("10. Nonconforming Product (สินค้าไม่เป็นไปตามข้อกำหนด)", level=1)
add_para("When purchased products are found to be nonconforming, the following actions shall be taken:")

add_bullet("Identify and segregate the nonconforming material (quarantine area)")
add_bullet("Issue NCR (FR-QC-002) to document the nonconformance")
add_bullet("Determine disposition: Return to Supplier / Rework / Use As Is / Scrap")
add_bullet("For return: issue SCAR (FR-PUR-005) and request corrective action from supplier")
add_bullet("Track supplier performance in Scorecard (FR-PUR-004)")

# ═══════════════════════════════════════════════════════════
# 11. RECORDS
# ═══════════════════════════════════════════════════════════
doc.add_heading("11. Records (บันทึกคุณภาพ)", level=1)
make_table(
    ["Record ID", "Title", "Retention", "ISO Clause"],
    [
        ["FR-PR-001", "Purchase Requisition", "3 years", "7.5.3"],
        ["FR-PUR-001", "Purchase Order", "3 years", "8.4.2 / 7.5.3"],
        ["FR-PUR-002", "Supplier Evaluation Form", "5 years", "8.4.1"],
        ["FR-PUR-003", "Approved Supplier List (ASL)", "Current", "8.4.1"],
        ["FR-QC-001", "Inspection Report", "5 years", "8.4.3 / 8.7"],
        ["FR-QC-002", "Nonconformance Report (NCR)", "5 years", "8.7 / 10.2"],
        ["FR-PUR-004", "Supplier Scorecard", "5 years", "8.4.1 / 9.1"],
        ["FR-PUR-005", "Supplier Corrective Action Request (SCAR)", "5 years", "10.2"],
    ],
    col_widths=[3.0, 5.5, 2.5, 3.0]
)

# ═══════════════════════════════════════════════════════════
# 12. REVISION HISTORY
# ═══════════════════════════════════════════════════════════
doc.add_heading("12. Revision History (ประวัติการแก้ไข)", level=1)
make_table(
    ["Rev.", "Date", "Description of Change", "Prepared", "Approved"],
    [
        ["00", "—", "Initial draft", "—", "—"],
        ["01", date.today().strftime("%d/%m/%Y"), "First issue (DRAFT)", "Purchasing Dept.", "—"],
    ],
    col_widths=[1.5, 3.0, 6.0, 3.0, 3.0]
)

# ═══════════════════════════════════════════════════════════
# APPROVAL
# ═══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_heading("Approval (การอนุมัติ)", level=2)
add_para("")
make_table(
    ["Role", "Name", "Signature", "Date"],
    [
        ["Prepared by:", "______________________", "______________________", "____/____/______"],
        ["Reviewed by:", "______________________", "______________________", "____/____/______"],
        ["Approved by:", "______________________", "______________________", "____/____/______"],
    ],
    col_widths=[3.0, 5.0, 5.0, 3.0]
)

doc.add_paragraph()
add_para("This document is confidential and intended for internal use only.",
         italic=True, size=8, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Unauthorized reproduction or distribution is prohibited.",
         italic=True, size=8, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──
output = r"C:\Working_Space\OpenWork\[DRAFT] QP-PUR-001 - Purchasing Procedure.docx"
doc.save(output)
print("OK")
