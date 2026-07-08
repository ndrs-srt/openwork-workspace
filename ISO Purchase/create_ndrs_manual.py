from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import date

doc = Document()

# ── Page setup: Portrait A4 ──
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)

# ── Colors ──
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
MED_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE_HEX = "D6E4F0"
RED       = RGBColor(0xC0, 0x39, 0x2B)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY      = RGBColor(0x66, 0x66, 0x66)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.space_before = Pt(1)
style.paragraph_format.line_spacing = 1.15

for lvl, sz, clr in [(1, 16, DARK_BLUE), (2, 13, MED_BLUE), (3, 11, DARK_BLUE)]:
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = 'Calibri'
    hs.font.size = Pt(sz)
    hs.font.color.rgb = clr
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(10 if lvl == 1 else 6)
    hs.paragraph_format.space_after  = Pt(4)


def add_para(text, bold=False, italic=False, size=10, align=None, color=None, space_after=3):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * level)
    return p


def shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=10, align=None, color=None, font_name='Calibri'):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.font.name = font_name
    r.font.size = Pt(size)
    r.bold = bold
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)


def make_table(headers, data, col_widths=None):
    ncols = len(headers)
    nrows = len(data) + 1
    t = doc.add_table(rows=nrows, cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell(c, h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(c, "2E75B6")
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            c = t.rows[ri+1].cells[ci]
            set_cell(c, str(val), size=9)
            if ri % 2 == 1:
                shade_cell(c, "F2F2F2")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# ═══════════════════════════════════════════════════════════
# NDRS-STYLE HEADER TABLE
# ═══════════════════════════════════════════════════════════
def add_doc_header(title_th="คู่มือการจัดซื้อ", title_en="Purchasing Manual",
                   doc_id="PM-001", rev="01", lang="TH/EN"):
    t = doc.add_table(rows=2, cols=4)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: dark blue bar
    texts_r0 = [
        f"NDRS-{doc_id}",
        f"{title_th}  {title_en}",
        f"Rev. {rev}",
        lang
    ]
    for i, txt in enumerate(texts_r0):
        c = t.rows[0].cells[i]
        set_cell(c, txt, bold=True, size=9, color=WHITE,
                 align=WD_ALIGN_PARAGRAPH.CENTER if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT)
        shade_cell(c, "1F3864")

    # Row 1: light blue info bar
    set_cell(t.rows[1].cells[0], f"Effective Date: {date.today().strftime('%d/%m/%Y')}",
             bold=True, size=8, color=RGBColor(0x1F, 0x38, 0x64))
    shade_cell(t.rows[1].cells[0], LIGHT_BLUE_HEX)
    set_cell(t.rows[1].cells[1], "Document Type: Procedure / คู่มือ", size=8)
    shade_cell(t.rows[1].cells[1], LIGHT_BLUE_HEX)
    set_cell(t.rows[1].cells[2], "Status: DRAFT", bold=True, size=8,
             color=RED, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(t.rows[1].cells[2], "FFF2CC")
    today_str = date.today().strftime('%d/%m/%Y')
    set_cell(t.rows[1].cells[3], f"Issued: {today_str}", size=8,
             align=WD_ALIGN_PARAGRAPH.RIGHT)
    shade_cell(t.rows[1].cells[3], LIGHT_BLUE_HEX)

    t.rows[0].cells[0].width = Cm(3.0)
    t.rows[0].cells[1].width = Cm(6.0)
    t.rows[0].cells[2].width = Cm(2.0)
    t.rows[0].cells[3].width = Cm(2.5)

    doc.add_paragraph()

def add_section_header(text):
    """Add a colored section header bar."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    set_cell(c, text, bold=True, size=11, color=WHITE)
    shade_cell(c, "2E75B6")
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para("NDRS", bold=True, size=28, color=DARK_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("QUALITY MANAGEMENT SYSTEM", bold=True, size=14, color=MED_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_para("คู่มือการจัดซื้อ", bold=True, size=22, color=DARK_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Purchasing Manual", bold=False, size=16, color=MED_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
add_para("ISO 9001:2015 — Clause 8.4", size=12, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_para(f"Document No.: NDRS-PM-001     |     Revision: 01", size=10, color=DARK_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f"Effective Date: {date.today().strftime('%d/%m/%Y')}     |     Status: DRAFT", size=10, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# DOCUMENT CONTROL (page header for all pages)
# ═══════════════════════════════════════════════════════════
add_doc_header("คู่มือการจัดซื้อ", "Purchasing Manual", "PM-001", "01", "TH/EN")

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════
doc.add_heading("สารบัญ (Table of Contents)", level=1)

toc_items = [
    ("1", "วัตถุประสงค์ (Purpose / Objective)"),
    ("2", "ขอบเขต (Scope)"),
    ("3", "คำจำกัดความ (Definitions)"),
    ("4", "หน้าที่ความรับผิดชอบ (Responsibilities)"),
    ("5", "ขั้นตอนการดำเนินงาน (Procedure)"),
    ("6", "เอกสารอ้างอิง (Reference Documents)"),
    ("7", "แบบฟอร์มที่ใช้ (Forms)"),
    ("8", "ประวัติการแก้ไข (Revision History)"),
    ("9", "ภาคผนวก (Appendix) — รายละเอียดเพิ่มเติม"),
]
for num, title in toc_items:
    add_bullet(f"ข้อ {num}: {title}")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. PURPOSE
# ═══════════════════════════════════════════════════════════
doc.add_heading("1. วัตถุประสงค์ (Purpose / Objective)", level=1)

add_para("คู่มือการจัดซื้อนี้จัดทำขึ้นโดยมีวัตถุประสงค์เพื่อกำหนดหลักเกณฑ์และขั้นตอนการจัดซื้อจัดจ้าง "
         "ของบริษัทฯ ให้เป็นมาตรฐานเดียวกัน สอดคล้องกับข้อกำหนด ISO 9001:2015 ข้อ 8.4 "
         "(การควบคุมผลิตภัณฑ์และบริการที่จัดหาจากภายนอก)")

purposes = [
    "เพื่อกำหนดแนวทางและขั้นตอนการดำเนินงานจัดซื้อให้เป็นมาตรฐาน",
    "เพื่อให้มั่นใจว่าผลิตภัณฑ์และบริการที่จัดซื้อมีคุณภาพตรงตามข้อกำหนดที่องค์กรต้องการ",
    "เพื่อคัดเลือก ประเมิน และติดตามผลผู้ขาย/ผู้รับเหมาช่วงอย่างมีประสิทธิภาพ",
    "เพื่อสร้างความโปร่งใส ตรวจสอบได้ และลดความเสี่ยงในการจัดซื้อ",
]
for p in purposes:
    add_bullet(p)

# ═══════════════════════════════════════════════════════════
# 2. SCOPE
# ═══════════════════════════════════════════════════════════
doc.add_heading("2. ขอบเขต (Scope)", level=1)

add_para("คู่มือนี้ครอบคลุมทุกกิจกรรมที่เกี่ยวข้องกับการจัดซื้อจัดจ้างวัสดุและบริการที่สำคัญ "
         "(Critical materials or services) ที่ส่งผลต่อคุณภาพของผลิตภัณฑ์ของบริษัทฯ")

scopes = [
    "วัตถุดิบและวัสดุสิ้นเปลืองที่ใช้ในการผลิต",
    "บรรจุภัณฑ์ที่สัมผัสกับผลิตภัณฑ์",
    "เครื่องจักรและอะไหล่ที่ส่งผลต่อคุณภาพผลิตภัณฑ์",
    "บริการภายนอก (Outsourced Services) เช่น การสอบเทียบ การทดสอบ การบำรุงรักษา",
    "กระบวนการผลิตที่จ้างช่วง (Subcontracted Manufacturing)",
]
for s in scopes:
    add_bullet(s)

add_para("ข้อยกเว้น: การจัดซื้อที่ไม่ส่งผลกระทบต่อคุณภาพของผลิตภัณฑ์ "
         "(เช่น อุปกรณ์สำนักงาน, วัสดุสิ้นเปลืองในฝ่ายบริหาร) ให้ใช้หลักเกณฑ์แบบย่อตามความเหมาะสม",
         italic=True, size=9)

# ═══════════════════════════════════════════════════════════
# 3. DEFINITIONS
# ═══════════════════════════════════════════════════════════
doc.add_heading("3. คำจำกัดความ (Definitions)", level=1)

doc.add_heading("3.1 เอกสารหลักในกระบวนการ", level=2)

defs_docs = [
    ("ใบขอซื้อ (PR - Purchase Requisition)",
     "เอกสารภายในที่จัดทำโดยผู้ขอซื้อ เพื่อแจ้งรายละเอียดความต้องการและขออนุมัติการจัดซื้อ"),
    ("ใบสั่งซื้อ (PO - Purchase Order)",
     "เอกสารทางการที่ออกให้แก่ผู้ขาย เพื่อยืนยันการสั่งซื้อสินค้าหรือบริการตามเงื่อนไขที่ตกลงกัน"),
    ("ใบกำกับภาษี / Invoice",
     "เอกสารที่ผู้ขายออกเพื่อเป็นหลักฐานการซื้อขายและเรียกเก็บเงิน"),
    ("ใบยืนยันการตรวจรับสินค้า (GRN)",
     "เอกสารภายในที่จัดทำโดยผู้รับสินค้า เพื่อยืนยันการรับมอบสินค้าที่ถูกต้อง"),
]
for title, desc in defs_docs:
    p = doc.add_paragraph()
    r = p.add_run(f"{title} - ")
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r = p.add_run(desc)
    r.font.name = 'Calibri'
    r.font.size = Pt(10)

doc.add_heading("3.2 บุคคลและกลุ่มที่เกี่ยวข้อง", level=2)

defs_people = [
    ("ผู้ขอซื้อ (Requester)", "พนักงานที่ต้องการสินค้าหรือบริการ และเริ่มต้นจัดทำ PR"),
    ("ผู้อนุมัติ (Approver)", "ผู้บังคับบัญชาที่มีอำนาจอนุมัติ PR / PO ตาม DOA"),
    ("ผู้รับผิดชอบการจัดซื้อ (Assigned Purchaser)", "ผู้ดำเนินการจัดซื้อหลังจากได้รับ PR ที่อนุมัติแล้ว"),
    ("ผู้ขาย / ผู้ให้บริการ (Supplier / Vendor)", "บุคคลภายนอกหรือนิติบุคคลที่จัดหาสินค้า/บริการ"),
    ("ผู้รับสินค้า (On-site Receiver)", "ผู้รับมอบสินค้าทางกายภาพและตรวจสอบเบื้องต้น"),
]
for title, desc in defs_people:
    p = doc.add_paragraph()
    r = p.add_run(f"{title} - ")
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r = p.add_run(desc)
    r.font.name = 'Calibri'
    r.font.size = Pt(10)

doc.add_heading("3.3 แนวคิดและเครื่องมือสำคัญ", level=2)

defs_concepts = [
    ("การประเมินผู้ขาย (Supplier Evaluation)", "กระบวนการพิจารณาคัดเลือกผู้ขายโดยใช้เกณฑ์ ราคา คุณภาพ การส่งมอบ บริการ และ EHS"),
    ("บัญชีรายชื่อผู้ขายฯ (ASL)", "รายชื่อผู้ขายที่ผ่านการประเมินและอนุมัติแล้ว"),
    ("การเทียบราคา (Quotation Comparison)", "การขอใบเสนอราคาจากผู้ขายหลายรายเพื่อเปรียบเทียบ"),
    ("3-Way Matching", "การตรวจสอบ PO, GRN, และ Invoice ให้ตรงกันก่อนจ่ายเงิน"),
    ("Lead Time", "ระยะเวลาตั้งแต่ออก PO จนถึงรับสินค้า"),
]
for title, desc in defs_concepts:
    p = doc.add_paragraph()
    r = p.add_run(f"{title} - ")
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r = p.add_run(desc)
    r.font.name = 'Calibri'
    r.font.size = Pt(10)

doc.add_heading("3.4 สถานะและเงื่อนไข", level=2)

defs_status = [
    ("วัสดุและบริการที่สำคัญ (Critical Materials)", "วัสดุ/บริการที่มีผลกระทบโดยตรงต่อคุณภาพผลิตภัณฑ์"),
    ("สถานะผู้ขาย (Supplier Status)", "Approved / Conditional / Disapproved"),
    ("เงื่อนไขการชำระเงิน (Payment Term)", "ข้อตกลงระยะเวลาชำระเงิน เช่น 30 วัน 60 วัน"),
]
for title, desc in defs_status:
    p = doc.add_paragraph()
    r = p.add_run(f"{title} - ")
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r = p.add_run(desc)
    r.font.name = 'Calibri'
    r.font.size = Pt(10)

# ═══════════════════════════════════════════════════════════
# 4. RESPONSIBILITIES
# ═══════════════════════════════════════════════════════════
doc.add_heading("4. หน้าที่ความรับผิดชอบ (Responsibilities)", level=1)

make_table(
    ["บทบาท / Role", "หน้าที่ความรับผิดชอบ (Responsibilities)"],
    [
        ["ผู้ขอซื้อ (Requester)",
         "จัดทำ PR ระบุสเปกและเหตุผล, เสนอผู้อนุมัติ, ยืนยันความถูกต้องทางเทคนิคเมื่อรับสินค้า"],
        ["ผู้อนุมัติ (Approver)",
         "อนุมัติ PR และ PO ตาม DOA, พิจารณาความจำเป็นและงบประมาณ"],
        ["ผู้รับผิดชอบการจัดซื้อ (Assigned Purchaser)",
         "คัดเลือกผู้ขาย, จัดทำ PO, ส่ง PO ให้ผู้ขาย/ผู้ขอซื้อ/ผู้รับสินค้า, รวบรวมเอกสารส่ง HAAP"],
        ["ผู้รับสินค้า (On-site Receiver)",
         "รับสินค้า, ตรวจสอบเบื้องต้น, แจ้งผู้ขอซื้อ, จัดทำ GRN, ส่งเอกสารให้ผู้รับผิดชอบการจัดซื้อ"],
        ["ทีม HAAP (Accounting/Finance)",
         "ตรวจสอบ 3-Way Matching, ดำเนินการชำระเงิน, เบิกจ่ายคืนพนักงาน"],
    ],
    col_widths=[4.0, 11.5]
)

# ═══════════════════════════════════════════════════════════
# 5. PROCEDURE
# ═══════════════════════════════════════════════════════════
doc.add_heading("5. ขั้นตอนการดำเนินงาน (Procedure)", level=1)

doc.add_heading("5.1 การขอซื้อและการอนุมัติ (Requisition and Approval)", level=2)

add_para("5.1.1 การจัดทำใบขอซื้อ (PR)", bold=True, size=10)
add_para("เมื่อพนักงาน (ผู้ขอซื้อ) มีความต้องการวัสดุหรือบริการที่สำคัญ ให้เริ่มต้นกระบวนการ "
         "โดยการจัดทำใบขอซื้อ (PR) ตามแบบฟอร์ม NDRS-FM-PUR-001 พร้อมระบุ:")
add_bullet("วันที่ต้องการใช้งานจริง")
add_bullet("คุณลักษณะเฉพาะ (Specification) ของสินค้า/บริการอย่างละเอียด")
add_bullet("วัตถุประสงค์ในการนำไปใช้งาน")
add_bullet("ผู้ขายที่แนะนำ (ถ้ามี)")
add_bullet("เอกสารแนบ เช่น ใบเสนอราคา, รูปภาพ, Link สินค้า")

add_para("5.1.2 การนำเสนอเพื่อขออนุมัติ", bold=True, size=10)
add_para("ผู้ขอซื้อนำส่ง PR ที่กรอกข้อมูลครบถ้วนแล้ว พร้อมเอกสารแนบ ให้ผู้อนุมัติพิจารณา")

add_para("5.1.3 การพิจารณาและอนุมัติ", bold=True, size=10)
add_para("ผู้อนุมัติตรวจสอบ PR ในประเด็นความจำเป็น ความสมเหตุสมผล และงบประมาณ")
add_bullet("กรณีอนุมัติ: ลงนามหรืออนุมัติผ่านช่องทางที่ตรวจสอบย้อนกลับได้")
add_bullet("กรณีไม่อนุมัติ: แจ้งเหตุผลและส่ง PR คืนผู้ขอซื้อ")

add_para("5.1.4 การส่งมอบเรื่อง", bold=True, size=10)
add_para("หลังจากได้รับอนุมัติแล้ว ผู้ขอซื้อส่งมอบ PR ที่อนุมัติให้ผู้รับผิดชอบการจัดซื้อดำเนินการ")

doc.add_heading("5.2 การดำเนินการจัดซื้อและออก PO (Purchasing and PO Issuance)", level=2)

add_para("5.2.1 การจำแนกประเภทผู้ขาย", bold=True, size=10)
add_para("ผู้รับผิดชอบการจัดซื้อจำแนกผู้ขายออกเป็น 3 ประเภท:")

make_table(
    ["ประเภท", "คำอธิบาย", "เอกสาร", "วิธีจัดซื้อ"],
    [
        ["ประเภทที่ 1\nStrategic Supplier",
         "ผู้จัดจำหน่ายรายใหญ่ระดับสากล\n(เช่น DigiKey, Mouser)", "Commercial Invoice",
         "ออก PO มาตรฐาน\nชำระผ่านบัตรเครดิต/โอน"],
        ["ประเภทที่ 2\nDomestic Supplier",
         "นิติบุคคลในประเทศที่ออก\nใบกำกับภาษีเต็มรูปแบบได้", "ใบกำกับภาษีเต็มรูปแบบ",
         "ออก PO มาตรฐาน\nเครดิตเทอม/โอน"],
        ["ประเภทที่ 3\nExceptional Case",
         "ร้านค้ารายย่อย/แพลตฟอร์มออนไลน์\n(Shopee, AliExpress)", "ไม่มีเอกสารภาษี",
         "สร้าง PO ภายใน (ห้ามส่ง)\nพนักงานสำรองจ่าย"],
    ],
    col_widths=[3.0, 4.5, 3.5, 4.5]
)

add_para("5.2.2 การจัดทำใบสั่งซื้อ (PO)", bold=True, size=10)
add_para("ผู้รับผิดชอบการจัดซื้อดำเนินการ:")
add_bullet("สร้าง PO ในระบบ PEAK ตามประเภทผู้ขาย")
add_bullet("นำส่ง PO ให้ผู้อนุมัติตรวจสอบและอนุมัติขั้นสุดท้าย")
add_bullet("ส่ง PO ที่อนุมัติแล้วให้: ผู้ขาย, ผู้ขอซื้อ, และผู้รับสินค้า")
add_bullet("กรณีประเภทที่ 3: สร้าง PO ภายในระบบเท่านั้น ห้ามส่งให้ผู้ขาย")

add_para("5.2.3 การยืนยันคำสั่งซื้อและการติดตาม", bold=True, size=10)
add_bullet("ขอใบยืนยันคำสั่งซื้อ (Order Confirmation) จากผู้ขาย")
add_bullet("ขอหมายเลขติดตามพัสดุ (Tracking Number) เมื่อผู้ขายจัดส่ง")

add_para("5.2.4 การจัดการที่อยู่จัดส่งและที่อยู่ออกบิล", bold=True, size=10)
add_bullet("ระบุ Bill To และ Ship To ให้ชัดเจนใน PO")
add_bullet("ยืนยันกับผู้ขายเพื่อให้ข้อมูลถูกต้องแก่บริษัทขนส่ง")

add_para("5.2.5 การจัดการชำระเงินล่วงหน้า", bold=True, size=10)
add_para("ในกรณีที่ผู้ขายกำหนดให้ชำระก่อนจัดส่ง ให้ขอ Proforma Invoice "
         "และจัดทำใบขอเบิกเงินจ่ายล่วงหน้าเสนอผู้อนุมัติ ก่อนส่งให้ทีม HAAP")

add_para("5.2.6 การจัดการสถานการณ์พิเศษ", bold=True, size=10)
add_bullet("การเปลี่ยนแปลง/ยกเลิก PO: ต้องได้รับการอนุมัติก่อนดำเนินการ")
add_bullet("Backorder: แจ้งผู้ขอซื้อเพื่อพิจารณาหรือหารือแนวทางแก้ไข")

doc.add_heading("5.3 การตรวจรับสินค้าและบริการ (Receiving of Goods and Services)", level=2)

add_para("5.3.1 การเตรียมการตรวจรับ", bold=True, size=10)
add_para("ผู้รับสินค้าและผู้ขอซื้อได้รับสำเนา PO จากผู้รับผิดชอบการจัดซื้อแล้ว")

add_para("5.3.2 การรับมอบสินค้าทางกายภาพ", bold=True, size=10)
add_bullet("ตรวจสอบสภาพภายนอกของบรรจุภัณฑ์")
add_bullet("ตรวจนับจำนวนให้ตรงกับเอกสารจัดส่ง")
add_bullet("ลงนามในเอกสารการจัดส่ง (ระบุหมายเหตุหากพบความเสียหาย)")

add_para("5.3.3 การแจ้งเตือนเพื่อทวนสอบ", bold=True, size=10)
add_para("ผู้รับสินค้าแจ้งผู้ขอซื้อทันที พร้อมส่งรูปถ่ายสินค้า ป้ายรุ่น และหมายเลขชิ้นส่วน")

add_para("5.3.4 การยืนยันความถูกต้อง", bold=True, size=10)
add_para("ผู้ขอซื้อยืนยันความถูกต้องทางเทคนิคโดยเปรียบเทียบรูปภาพกับสเปกใน PR")

add_para("5.3.5 การจัดทำเอกสารยืนยันการตรวจรับ", bold=True, size=10)
add_bullet("ผู้รับสินค้าจัดทำและลงนามในเอกสารยืนยันการตรวจรับ (NDRS-FM-PUR-003)")
add_bullet("กรณีสินค้าไม่ถูกต้อง ให้หมายเหตุและแจ้งผู้รับผิดชอบการจัดซื้อ")

add_para("5.3.6 การส่งมอบเอกสาร", bold=True, size=10)
add_para("ผู้รับสินค้าส่งเอกสารยืนยันการตรวจรับที่ลงนามแล้วให้ผู้รับผิดชอบการจัดซื้อ")

doc.add_heading("5.4 การรวบรวมเอกสารส่งมอบให้ทีม HAAP (Document Consolidation)", level=2)

add_para("5.4.1 การทวนสอบความครบถ้วน", bold=True, size=10)
add_bullet("สำหรับประเภท 1 และ 2: PO + Invoice/ใบกำกับภาษี + เอกสารยืนยันการตรวจรับ")
add_bullet("สำหรับประเภท 3 (สำรองจ่าย): PO ภายใน + ชุดเอกสารเบิกคืนตามภาคผนวก ก")

add_para("5.4.2 การส่งมอบเรื่องให้ทีม HAAP", bold=True, size=10)
add_para("เมื่อเอกสารครบถ้วน ผู้รับผิดชอบการจัดซื้อส่งมอบเรื่องให้ทีม HAAP เพื่อดำเนินการตรวจสอบ "
         "3-Way Matching และชำระเงิน")

add_para("5.4.3 การสิ้นสุดกระบวนการ", bold=True, size=10)
add_para("ณ จุดนี้ หน้าที่ความรับผิดชอบตามระเบียบปฏิบัติการจัดซื้อฉบับนี้สิ้นสุดลง")

# ═══════════════════════════════════════════════════════════
# 6. REFERENCE DOCUMENTS
# ═══════════════════════════════════════════════════════════
doc.add_heading("6. เอกสารอ้างอิง (Reference Documents)", level=1)

make_table(
    ["รหัสเอกสาร", "ชื่อเอกสาร"],
    [
        ["QM-001", "Quality Manual / คู่มือคุณภาพ"],
        ["NDRS-WI-PUR-001", "วิธีการจัดทำใบขอซื้อ (PR) / PR Creation Guide"],
        ["NDRS-WI-PUR-002", "วิธีการสร้างใบสั่งซื้อ (PO) ในโปรแกรม PEAK"],
        ["NDRS-WI-PUR-003", "วิธีการจัดการที่อยู่จัดส่งสำหรับผู้ให้บริการขนส่ง"],
        ["NDRS-WI-PUR-004", "คู่มือและเช็คลิสต์สำหรับประเมินและคัดเลือกผู้ขายรายใหม่"],
        ["NDRS-WI-PUR-005", "บัญชีรายชื่อผู้ขายฯ (ASL) / Approved Supplier List"],
        ["NDRS-WI-PUR-006", "วิธีการเบิกเงินคืนกรณีพนักงานสำรองจ่าย"],
        ["WI-QC-001", "ขั้นตอนการตรวจสอบขาเข้า (Incoming Inspection)"],
        ["ISO 9001:2015", "Quality Management Systems — Requirements"],
    ],
    col_widths=[4.5, 11.0]
)

# ═══════════════════════════════════════════════════════════
# 7. FORMS
# ═══════════════════════════════════════════════════════════
doc.add_heading("7. แบบฟอร์มที่ใช้ (Forms)", level=1)

add_para("แบบฟอร์มหลักที่ใช้ในกระบวนการจัดซื้อ มีดังนี้:")

make_table(
    ["รหัสเอกสาร (Form ID)", "ชื่อแบบฟอร์ม (Form Name)", "วัตถุประสงค์"],
    [
        ["NDRS-FM-PUR-001", "ใบขอซื้อ (Purchase Requisition Form)",
         "ใช้สำหรับพนักงานทุกคนในการแจ้งความต้องการและขออนุมัติการจัดซื้อ"],
        ["NDRS-FM-PUR-002", "ใบรับรองแทนใบเสร็จรับเงิน",
         "ใช้สำหรับพนักงานในการเบิกเงินคืน กรณีสำรองจ่ายสำหรับผู้ขายประเภทที่ 3"],
        ["NDRS-FM-PUR-003", "เอกสารยืนยันการตรวจรับสินค้า",
         "ใช้สำหรับผู้รับสินค้าเพื่อยืนยันการรับมอบสินค้าที่ถูกต้อง"],
        ["NDRS-FM-PUR-004", "ใบประเมินผู้ขายรายใหม่",
         "ใช้สำหรับบันทึกผลการประเมินผู้ขายรายใหม่ตามเกณฑ์ที่กำหนด"],
        ["(สร้างจาก PEAK)", "ใบสั่งซื้อ (Purchase Order)",
         "เอกสารทางการที่สร้างจากระบบ PEAK เพื่อส่งให้ผู้ขาย"],
        ["(ไฟล์ควบคุม)", "บัญชีรายชื่อผู้ขายฯ (ASL)",
         "Spreadsheet ที่ใช้ควบคุมรายชื่อและสถานะผู้ขาย"],
    ],
    col_widths=[4.5, 5.5, 6.5]
)

# ═══════════════════════════════════════════════════════════
# 8. REVISION HISTORY
# ═══════════════════════════════════════════════════════════
doc.add_heading("8. ประวัติการแก้ไข (Revision History)", level=1)

make_table(
    ["Rev.", "Date", "Description of Change", "Prepared", "Approved"],
    [
        ["00", "—", "Initial draft", "—", "—"],
        ["01", date.today().strftime("%d/%m/%Y"), "First issue (DRAFT)", "Purchasing Dept.", "—"],
    ],
    col_widths=[1.5, 3.0, 6.0, 3.0, 3.0]
)

# ═══════════════════════════════════════════════════════════
# 9. APPENDIX
# ═══════════════════════════════════════════════════════════
doc.add_heading("9. ภาคผนวก (Appendix)", level=1)

# ── Appendix A ──
doc.add_heading("ภาคผนวก ก — นโยบายพนักงานสำรองจ่ายสำหรับผู้ให้บริการกรณีพิเศษ (ประเภทที่ 3)", level=2)

add_para("1. หลักการและเงื่อนไข", bold=True, size=10)
add_bullet("นโยบายนี้ใช้เป็นข้อยกเว้นในกรณีจำเป็นเร่งด่วนที่ต้องจัดซื้อจากผู้ให้บริการกรณีพิเศษ "
           "(ประเภทที่ 3) ซึ่งไม่สามารถออกเอกสารทางภาษีได้")
add_bullet("วงเงินจำกัดไม่เกิน [ระบุจำนวน] บาทต่อครั้ง")
add_bullet("ผู้อนุมัติต้องรับทราบว่าบริษัทฯ ไม่สามารถใช้สิทธิ์ทางภาษีซื้อ (Input VAT) ได้")

add_para("2. ขั้นตอนการปฏิบัติและการเบิกเงินคืน", bold=True, size=10)
add_bullet("2.1 พนักงานสำรองจ่ายด้วยเงินส่วนตัวไปก่อน")
add_bullet("2.2 รวบรวมหลักฐาน: ภาพหน้าจอคำสั่งซื้อ, หลักฐานการชำระเงิน, ภาพถ่ายสินค้า, "
           "และใบรับรองแทนใบเสร็จรับเงิน (NDRS-FM-PUR-002)")
add_bullet("2.3 ขออนุมัติจากผู้อนุมัติ (หัวหน้างาน)")
add_bullet("2.4 ส่งเอกสารให้ผู้รับผิดชอบการจัดซื้อเพื่อรวบรวมส่งทีม HAAP")

doc.add_paragraph()
# ── Appendix B ──
doc.add_heading("ภาคผนวก ข — ตารางอำนาจอนุมัติการจัดซื้อ (DOA)", level=2)

add_para("ตารางอำนาจอนุมัติ (Delegation of Authority) สำหรับการอนุมัติ PR และ PO:")

make_table(
    ["ระดับ (Level)", "ตำแหน่ง (Position)", "วงเงินอนุมัติต่อครั้ง (บาท)"],
    [
        ["1", "Team Leader / หัวหน้าทีม", "\u2264 20,000"],
        ["2", "ผู้อำนวยการฝ่าย / หัวหน้า BU", "20,001 - 200,000"],
        ["3", "กรรมการผู้จัดการ (Managing Director)", "> 200,000"],
    ],
    col_widths=[3.0, 6.5, 6.0]
)

add_para("หมายเหตุ:", italic=True, size=9, color=GRAY)
add_para("ตารางอำนาจอนุมัตินี้ใช้บังคับทั้งการอนุมัติ PR และ PO", italic=True, size=9, color=GRAY)
add_para("การขอซื้อที่มีมูลค่าสูงกว่าวงเงินของตำแหน่งใด ให้เสนอขออนุมัติจากผู้บังคับบัญชาในระดับถัดไป",
         italic=True, size=9, color=GRAY)

doc.add_paragraph()

# ── Appendix C ──
doc.add_heading("ภาคผนวก ค — สรุปแนวทางการจัดการผู้ให้บริการภายนอก", level=2)

make_table(
    ["ประเภทผู้ให้บริการ", "เอกสารหลักที่ต้องได้รับ", "วิธีการจัดซื้อ", "การชำระเงิน"],
    [
        ["ประเภทที่ 1\nStrategic Supplier\n(เช่น DigiKey, Mouser)",
         "Commercial Invoice", "ออก PO มาตรฐาน", "บัตรเครดิตบริษัท,\nโอนเงิน"],
        ["ประเภทที่ 2\nDomestic Supplier\n(บริษัทในประเทศ)",
         "ใบกำกับภาษีเต็มรูปแบบ", "ออก PO มาตรฐาน", "เครดิตเทอม,\nโอนเงิน"],
        ["ประเภทที่ 3\nExceptional Case\n(Shopee, AliExpress)",
         "ไม่มีเอกสารภาษี", "สร้าง PO ภายใน\n(ห้ามส่งผู้ขาย)", "พนักงานสำรองจ่าย\nเท่านั้น"],
    ],
    col_widths=[4.0, 4.0, 4.0, 3.5]
)

# ── END ──
doc.add_paragraph()
add_para("— จบเอกสาร (End of Document) —", bold=True, size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
add_para(f"จัดทำโดย: ฝ่ายจัดซื้อ (Purchasing Dept.)  |  "
         f"วันที่: {date.today().strftime('%d/%m/%Y')}  |  "
         f"เอกสารควบคุม: NDRS-PM-001 ฉบับที่ 01",
         size=8, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──
output = r"C:\Working_Space\OpenWork\คู่มือการสั่งซื้อ_ISO_Purchasing_Manual_NDRS.docx"
doc.save(output)
print("OK")
