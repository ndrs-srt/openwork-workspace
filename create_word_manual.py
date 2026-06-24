"""Generate ISO Purchasing Manual as MS Word (.docx)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import date
import os

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.page_width = Cm(29.7)   # A4 landscape
section.page_height = Cm(21.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# ── Style definitions ──
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
MED_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

# Heading styles
for lvl, sz, clr in [(1, 18, DARK_BLUE), (2, 14, DARK_BLUE), (3, 12, MED_BLUE)]:
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = 'Calibri'
    hs.font.size = Pt(sz)
    hs.font.color.rgb = clr
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
    hs.paragraph_format.space_after = Pt(6)


def add_para(text, bold=False, italic=False, size=10, align=None, color=None, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p


def shade_cell(cell, color_hex):
    """Apply background color to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=10, align=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)


def make_table(headers, data, col_widths=None):
    """Create a formatted table."""
    ncols = len(headers)
    nrows = len(data) + 1
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, "2E75B6")

    # Data rows
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            set_cell(cell, str(val), size=10)
            if ri % 2 == 1:
                shade_cell(cell, "F2F2F2")

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    else:
        for row in table.rows:
            for i in range(ncols):
                row.cells[i].width = Cm(14.0 / ncols)  # landscape ~14cm usable

    doc.add_paragraph()  # spacer
    return table


# =====================================================================
# COVER PAGE
# =====================================================================
for _ in range(4):
    doc.add_paragraph()

add_para("คู่มือการจัดซื้อ", bold=True, size=26, color=DARK_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Purchasing Manual", bold=True, size=18, color=MED_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_para("ISO 9001:2015 — ข้อกำหนด 8.4", size=14, color=MED_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

add_para("เอกสารหมายเลข: PM-001", bold=True, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f"ฉบับที่: 01  |  วันที่ประกาศใช้: {date.today().strftime('%d/%m/%Y')}",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("แผนก: ฝ่ายจัดซื้อ (Purchasing Department)", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS (Manual)
# =====================================================================
doc.add_heading("สารบัญ (Table of Contents)", level=1)

toc_items = [
    ("1", "วัตถุประสงค์ (Objective)"),
    ("2", "ขอบเขต (Scope)"),
    ("3", "เอกสารอ้างอิง (Reference Documents)"),
    ("4", "คำจำกัดความ (Definitions)"),
    ("5", "ความรับผิดชอบ (Responsibilities)"),
    ("6", "ขั้นตอนการจัดซื้อ (Purchasing Process)"),
    ("7", "การประเมินและคัดเลือกผู้ขาย (Supplier Evaluation)"),
    ("8", "ใบสั่งซื้อ / ข้อตกลงการจัดซื้อ (Purchase Order)"),
    ("9", "การตรวจสอบสินค้าที่ได้รับ (Incoming Inspection)"),
    ("10", "บันทึกคุณภาพ (Quality Records)"),
    ("11", "การแก้ไขและการปรับปรุง (Corrective Actions)"),
]
for num, title in toc_items:
    add_bullet(f"ข้อ {num}: {title}")

doc.add_page_break()

# =====================================================================
# 1. OBJECTIVE
# =====================================================================
doc.add_heading("1. วัตถุประสงค์ (Objective)", level=1)

add_para("คู่มือการจัดซื้อนี้จัดทำขึ้นโดยมีวัตถุประสงค์ดังนี้:", size=10)
objectives = [
    "เพื่อกำหนดแนวทางและขั้นตอนการดำเนินงานจัดซื้อให้เป็นมาตรฐาน สอดคล้องกับข้อกำหนด ISO 9001:2015 ข้อ 8.4 (การควบคุมผลิตภัณฑ์และบริการที่จัดหาจากภายนอก)",
    "เพื่อให้มั่นใจว่าผลิตภัณฑ์และบริการที่จัดซื้อมีคุณภาพตรงตามข้อกำหนดที่องค์กรต้องการ",
    "เพื่อคัดเลือก ประเมิน และติดตามผลผู้ขาย/ผู้รับเหมาช่วงอย่างมีประสิทธิภาพ",
    "เพื่อสร้างความโปร่งใส ตรวจสอบได้ และลดความเสี่ยงในการจัดซื้อ",
]
for obj in objectives:
    add_bullet(obj)

# =====================================================================
# 2. SCOPE
# =====================================================================
doc.add_heading("2. ขอบเขต (Scope)", level=1)

add_para("คู่มือนี้ครอบคลุมขอบเขตดังต่อไปนี้:", size=10)
scopes = [
    "ครอบคลุมกระบวนการจัดซื้อทั้งหมดขององค์กร ตั้งแต่การระบุความต้องการจนถึงการรับสินค้า/บริการ",
    "ครอบคลุมการจัดซื้อวัตถุดิบ วัสดุสิ้นเปลือง ครุภัณฑ์ และบริการภายนอกที่ส่งผลต่อคุณภาพของผลิตภัณฑ์",
    "ครอบคลุมการประเมิน คัดเลือก และการติดตามผลผู้ขาย (Supplier) และผู้รับเหมาช่วง (Subcontractor)",
    "ยกเว้น: การจัดซื้อที่ไม่ส่งผลกระทบต่อคุณภาพของผลิตภัณฑ์ (เช่น อุปกรณ์สำนักงานทั่วไป) ให้นำเฉพาะแนวปฏิบัติที่ดีมาใช้",
]
for s in scopes:
    add_bullet(s)

# =====================================================================
# 3. REFERENCE DOCUMENTS
# =====================================================================
doc.add_heading("3. เอกสารอ้างอิง (Reference Documents)", level=1)

make_table(
    ["ลำดับ", "รหัสเอกสาร", "ชื่อเอกสาร", "หมายเลข ISO"],
    [
        ["1", "QM-001", "คู่มือคุณภาพ (Quality Manual)", "ISO 9001:2015"],
        ["2", "WI-PUR-001", "ขั้นตอนการจัดซื้อ (Purchasing Procedure)", "8.4"],
        ["3", "WI-PUR-002", "ขั้นตอนการประเมินผู้ขาย (Supplier Evaluation)", "8.4.1"],
        ["4", "WI-QC-001", "ขั้นตอนการตรวจสอบขาเข้า (Incoming Inspection)", "8.4.3"],
        ["5", "FR-PUR-001", "แบบฟอร์มใบสั่งซื้อ (Purchase Order Form)", "8.4.2"],
        ["6", "FR-PUR-002", "แบบฟอร์มประเมินผู้ขาย (Supplier Evaluation Form)", "8.4.1"],
        ["7", "FR-PUR-003", "แบบฟอร์มทะเบียนผู้ขาย (Approved Supplier List)", "8.4.1"],
        ["8", "FR-QC-001", "รายงานตรวจสอบคุณภาพ (Inspection Report)", "8.4.3"],
        ["9", "—", "ข้อกำหนดระบบบริหารคุณภาพ ISO 9001:2015", "ทั้งหมด"],
        ["10", "—", "ข้อกำหนดของลูกค้า (Customer Specifications)", "7.2 / 8.2"],
    ],
    col_widths=[1.5, 3.0, 8.0, 3.0]
)

# =====================================================================
# 4. DEFINITIONS
# =====================================================================
doc.add_heading("4. คำจำกัดความ (Definitions)", level=1)

make_table(
    ["คำศัพท์", "คำอธิบาย", "เทียบเคียง ISO"],
    [
        ["ผู้ขาย / Supplier", "บุคคล หรือ นิติบุคคล ที่จัดหาผลิตภัณฑ์หรือบริการให้องค์กร", "External Provider (8.4)"],
        ["ผู้รับเหมาช่วง / Subcontractor", "ผู้ที่ได้รับมอบหมายให้ดำเนินการบางส่วนของกระบวนการผลิตแทนองค์กร", "Outsourced Process (8.4)"],
        ["ใบสั่งซื้อ / Purchase Order (PO)", "เอกสารที่ระบุรายละเอียดสินค้า ปริมาณ ราคา เงื่อนไข และกำหนดส่ง", "Purchase Order (8.4.2)"],
        ["ทะเบียนผู้ขาย / Approved Supplier List (ASL)", "รายชื่อผู้ขายที่ผ่านการประเมินและอนุมัติให้จัดซื้อได้", "Approved Supplier List (8.4.1)"],
        ["การประเมินผู้ขาย / Supplier Evaluation", "กระบวนการประเมินความสามารถของผู้ขายตามเกณฑ์ที่กำหนด", "Evaluation (8.4.1)"],
        ["สินค้าไม่เป็นไปตามข้อกำหนด / Nonconforming Product", "สินค้าหรือบริการที่ไม่ตรงตามข้อกำหนดที่ระบุใน PO หรือสเปก", "Nonconformity (8.7 / 10.2)"],
        ["ข้อกำหนดการจัดซื้อ / Purchase Requirement", "รายละเอียดคุณสมบัติ มาตรฐาน หรือสเปกที่ต้องการจากผู้ขาย", "Purchasing Requirements (8.4.2)"],
    ],
    col_widths=[3.5, 9.0, 3.5]
)

# =====================================================================
# 5. RESPONSIBILITIES
# =====================================================================
doc.add_heading("5. ความรับผิดชอบ (Responsibilities)", level=1)

make_table(
    ["บทบาท / ฝ่าย", "หน้าที่ความรับผิดชอบหลัก", "เอกสารที่เกี่ยวข้อง"],
    [
        ["ผู้บริหารสูงสุด (Top Management)", "อนุมัตินโยบายการจัดซื้อ อนุมัติผู้ขายรายใหม่ อนุมัติวงเงินสูง", "QM-001, FR-PUR-002"],
        ["ฝ่ายจัดซื้อ (Purchasing Dept.)", "ดำเนินการจัดซื้อตามขั้นตอน คัดเลือกผู้ขาย จัดทำ PO ติดตามงาน", "WI-PUR-001, FR-PUR-001"],
        ["ฝ่ายควบคุมคุณภาพ (QC Dept.)", "ตรวจสอบคุณภาพสินค้าขาเข้า อนุมัติ/ปฏิเสธสินค้า", "WI-QC-001, FR-QC-001"],
        ["ฝ่ายคลังสินค้า (Warehouse)", "รับสินค้า ตรวจทานจำนวน จัดเก็บ และแจ้ง QC ตรวจสอบ", "WI-WH-001"],
        ["ฝ่ายวิศวกรรม / ผลิต (Engineering)", "กำหนดสเปก มาตรฐาน และข้อกำหนดทางเทคนิคของสินค้า", "Specification Sheet"],
        ["ผู้ร้องขอซื้อ (Requestor)", "ระบุความต้องการ จัดทำ PR (Purchase Requisition)", "FR-PR-001"],
    ],
    col_widths=[4.0, 8.5, 3.5]
)

# =====================================================================
# 6. PURCHASING PROCESS
# =====================================================================
doc.add_heading("6. ขั้นตอนการจัดซื้อ (Purchasing Process)", level=1)

add_para("กระบวนการจัดซื้อขององค์กรประกอบด้วย 13 ขั้นตอนหลัก ดังนี้:", size=10)

make_table(
    ["ขั้นตอน", "ผู้รับผิดชอบ", "รายละเอียดการปฏิบัติงาน", "เอกสาร/บันทึก", "ข้อกำหนด ISO"],
    [
        ["1. ระบุความต้องการ", "ผู้ร้องขอซื้อ", "ระบุรายละเอียดสินค้า/บริการ ปริมาณ กำหนดส่ง และสเปก", "PR (FR-PR-001)", "8.4.2"],
        ["2. ขออนุมัติจัดซื้อ", "หัวหน้าฝ่าย", "ตรวจสอบความจำเป็น, งบประมาณ, และอนุมัติ PR", "PR อนุมัติ", "5.1 / 7.1"],
        ["3. คัดเลือกผู้ขาย", "ฝ่ายจัดซื้อ", "เลือกจาก ASL หรือประเมินผู้ขายรายใหม่", "ASL (FR-PUR-003)", "8.4.1"],
        ["4. ขอ quotation", "ฝ่ายจัดซื้อ", "ขอใบเสนอราคาอย่างน้อย 3 ราย (มูลค่าสูง)", "Quotation / RFQ", "8.4.2"],
        ["5. จัดทำ PO", "ฝ่ายจัดซื้อ", "ระบุสินค้า ปริมาณ ราคา วันที่ส่ง เงื่อนไขชัดเจน", "PO (FR-PUR-001)", "8.4.2"],
        ["6. อนุมัติ PO", "ผู้มีอำนาจ", "ตรวจสอบความถูกต้องก่อนส่งผู้ขาย", "PO อนุมัติ", "5.1"],
        ["7. สั่งซื้อ / ส่ง PO", "ฝ่ายจัดซื้อ", "ส่ง PO ทาง Email / Fax / ระบบ EDI", "PO ที่ส่ง", "8.4.2"],
        ["8. ติดตาม / expediting", "ฝ่ายจัดซื้อ", "ติดตามความคืบหน้า ยืนยันกำหนดส่ง", "Expediting Log", "8.4.2"],
        ["9. รับสินค้า", "คลังสินค้า", "นับจำนวน ตรวจสอบ Delivery Note เทียบ PO", "DN / GRN", "8.4.3"],
        ["10. ตรวจสอบคุณภาพ", "QC", "ตรวจตาม Inspection Plan", "Insp. Report (FR-QC-001)", "8.4.3"],
        ["11. รับเข้า / reject", "คลัง / QC", "OK → รับเข้าคลัง / NG → NCR", "GRN / NCR", "8.4.3 / 8.7"],
        ["12. วางบิล / จ่ายเงิน", "ฝ่ายบัญชี", "ตรวจสอบเอกสารครบถ้วน จ่ายเงินตามเงื่อนไข", "Invoice", "—"],
        ["13. บันทึก / วัดผล", "จัดซื้อ / QC", "บันทึกผลการส่งมอบ คุณภาพ Feedback", "Scorecard", "8.4.1 / 9.1"],
    ],
    col_widths=[3.0, 2.5, 6.0, 3.0, 2.5]
)

# =====================================================================
# 7. SUPPLIER EVALUATION
# =====================================================================
doc.add_heading("7. การประเมินและคัดเลือกผู้ขาย (Supplier Evaluation)", level=1)

doc.add_heading("7.1 เกณฑ์การประเมินผู้ขายรายใหม่", level=2)

make_table(
    ["เกณฑ์การประเมิน", "รายละเอียด", "น้ำหนัก (%)"],
    [
        ["คุณภาพ (Quality)", "ระบบคุณภาพ ISO 9001, การควบคุมกระบวนการ, ประวัติคุณภาพสินค้า", "35%"],
        ["ราคา (Price)", "ความสามารถในการแข่งขัน, ความโปร่งใสของโครงสร้างราคา", "20%"],
        ["การส่งมอบ (Delivery)", "ความสามารถส่งมอบตรงเวลา, ความยืดหยุ่น, Lead Time", "20%"],
        ["บริการ (Service)", "การตอบสนอง, การให้คำปรึกษาทางเทคนิค, บริการหลังการขาย", "10%"],
        ["ความน่าเชื่อถือ (Reliability)", "ประวัติบริษัท, ฐานะการเงิน, ข้อพิพาท, การอ้างอิงลูกค้า", "10%"],
        ["สิ่งแวดล้อม/ความปลอดภัย (EHS)", "การปฏิบัติตามกฎหมาย, ใบอนุญาต, ISO 14001 / ISO 45001", "5%"],
    ],
    col_widths=[4.5, 9.5, 2.5]
)

doc.add_heading("7.2 ระดับผลการประเมิน (Evaluation Rating)", level=2)

make_table(
    ["คะแนนรวม", "ระดับ", "การดำเนินการ"],
    [
        ["\u2265 85%", "A — ดีเยี่ยม (Excellent)", "ปรับปรุงต่อเนื่อง, พิจารณาสิทธิพิเศษ"],
        ["70% \u2013 84%", "B — ดี (Good)", "ติดตามผลเป็นระยะ, แจ้งจุดที่ควรปรับปรุง"],
        ["50% \u2013 69%", "C — พอใช้ (Fair)", "กำหนดแผนแก้ไข (CAP) และติดตามอย่างใกล้ชิด"],
        ["< 50%", "D — ต้องปรับปรุง (Poor)", "พิจารณาเพิกถอนจาก ASL / หาผู้ขายรายใหม่"],
    ],
    col_widths=[3.0, 5.0, 8.0]
)

doc.add_heading("7.3 ความถี่ในการประเมินซ้ำ (Re-evaluation Frequency)", level=2)

make_table(
    ["ระดับผู้ขาย", "ความถี่", "วิธี"],
    [
        ["A", "ทุก 12 เดือน", "ทบทวน Performance Scorecard + สอบถาม"],
        ["B", "ทุก 6 เดือน", "ทบทวน Performance Scorecard + ส่งแบบประเมิน"],
        ["C", "ทุก 3 เดือน", "ทบทวน Performance Scorecard + อาจตรวจสอบหน้างาน"],
        ["D", "ทันที", "ดำเนินการตาม CAP หรือถอดจาก ASL"],
    ],
    col_widths=[2.5, 3.5, 10.0]
)

# =====================================================================
# 8. PURCHASE ORDER
# =====================================================================
doc.add_heading("8. ใบสั่งซื้อ / ข้อตกลงการจัดซื้อ (Purchase Order)", level=1)

add_para("ใบสั่งซื้อ (Purchase Order: PO) เป็นเอกสารสำคัญที่ใช้ในการสั่งซื้อสินค้าหรือบริการจากผู้ขาย "
         "โดยต้องระบุรายละเอียดดังต่อไปนี้:", size=10)

make_table(
    ["รายการ", "คำอธิบาย"],
    [
        ["เลขที่ PO", "PO-XXXX-XXX (กำหนดโดยระบบ)"],
        ["วันที่ออก PO", "DD/MM/YYYY"],
        ["ชื่อผู้ขาย", "ชื่อบริษัท / นิติบุคคล"],
        ["ที่อยู่ผู้ขาย", "ที่อยู่จัดส่ง / ที่อยู่ออก Invoice"],
        ["เลขที่ผู้เสียภาษี", "เลขประจำตัวผู้เสียภาษี"],
        ["ผู้ติดต่อ", "ชื่อผู้ประสานงาน + เบอร์โทร / Email"],
        ["เงื่อนไขการชำระเงิน", "เงินสด / เครดิต 30 วัน / เครดิต 60 วัน / อื่นๆ"],
        ["เงื่อนไขการส่งมอบ", "Incoterm: EXW / FOB / CIF / DDP / อื่นๆ"],
        ["วันที่ต้องการสินค้า", "DD/MM/YYYY"],
    ],
    col_widths=[4.5, 11.5]
)

doc.add_heading("8.1 ข้อกำหนด ISO ที่ต้องระบุใน PO", level=2)

add_para("ตามข้อกำหนด ISO 9001:2015 ข้อ 8.4.2 ใบสั่งซื้อต้องมีรายการตรวจสอบดังนี้:", size=10)

make_table(
    ["ข้อ", "รายการตรวจสอบ"],
    [
        ["8.4.2 a)", "ข้อกำหนดสำหรับผลิตภัณฑ์/บริการที่ต้องการจัดซื้อ (Spec / Drawing / Standard)"],
        ["8.4.2 b)", "ข้อกำหนดสำหรับการอนุมัติ: ผลิตภัณฑ์, ขั้นตอน, กระบวนการ, อุปกรณ์"],
        ["8.4.2 c)", "ข้อกำหนดสำหรับบุคลากรที่มีคุณสมบัติ (ถ้าต้องการ)"],
        ["8.4.2 d)", "ข้อกำหนดของระบบบริหารคุณภาพ (เช่น ISO 9001)"],
    ],
    col_widths=[2.5, 13.0]
)

add_para("หมายเหตุ: ทุก PO ต้องมีลายเซ็นผู้อนุมัติก่อนส่งให้ผู้ขาย และต้องจัดเก็บ PO อย่างน้อย 3 ปี "
         "(ตามข้อกำหนด ISO 9001:2015 ข้อ 7.5.3)", italic=True, size=9, color=RGBColor(0x55, 0x55, 0x55))

# =====================================================================
# 9. INCOMING INSPECTION
# =====================================================================
doc.add_heading("9. การตรวจสอบสินค้าที่ได้รับ (Incoming Inspection)", level=1)

make_table(
    ["ขั้นตอน", "ผู้รับผิดชอบ", "รายละเอียด", "เอกสาร", "มาตรฐาน ISO"],
    [
        ["1. รับเอกสารส่งของ", "คลังสินค้า", "ตรวจสอบ Delivery Note, Invoice กับ PO", "DN / PO", "8.4.3"],
        ["2. ตรวจนับจำนวน", "คลังสินค้า", "นับจำนวนจริง เทียบกับเอกสาร", "GRN", "8.4.3"],
        ["3. ตรวจสอบสภาพ", "คลังสินค้า", "ตรวจสอบความเสียหายของบรรจุภัณฑ์", "GRN / Photo", "8.4.3"],
        ["4. แจ้ง QC", "คลังสินค้า", "ติดต่อ QC เพื่อตรวจสอบคุณภาพตามแผน", "Inspection Request", "8.4.3"],
        ["5. ตรวจสอบคุณภาพ", "QC", "ตรวจตาม Inspection Plan / Spec", "Insp. Report", "8.4.3"],
        ["6. ตัดสินใจ", "QC", "Accept / Reject / Conditional Accept", "Insp. Report", "8.4.3"],
        ["7. รับเข้า / คืน", "คลัง/QC", "Accept → รับเข้า / Reject → คืน + NCR", "GRN/NCR/SCAR", "8.4.3/8.7"],
    ],
    col_widths=[3.0, 2.5, 5.0, 3.5, 2.5]
)

doc.add_heading("9.1 ตัวอย่างแผนการตรวจสอบ (Sampling Inspection Plan)", level=2)

make_table(
    ["ประเภทสินค้า", "ระดับการตรวจ", "จำนวนตัวอย่าง", "เกณฑ์ยอมรับ", "วิธีตรวจ"],
    [
        ["วัตถุดิบ A (Critical)", "100%", "ทั้งหมด", "Zero Defect (C=0)", "ตาม Spec Sheet"],
        ["วัตถุดิบ B (Major)", "สุ่ม (Random)", "\u221AN หรือ MIL-STD-1916", "AQL 1.0%", "ตาม Spec Sheet"],
        ["วัตถุดิบ C (Minor)", "สุ่ม (Random)", "MIL-STD-1916 Level II", "AQL 4.0%", "Visual + Dimension"],
        ["บริการภายนอก", "—", "ตรวจ Cert / Report", "—", "เอกสารประกอบ"],
    ],
    col_widths=[4.0, 2.5, 3.5, 3.0, 3.0]
)

# =====================================================================
# 10. QUALITY RECORDS
# =====================================================================
doc.add_heading("10. บันทึกคุณภาพ (Quality Records)", level=1)

add_para("บันทึกคุณภาพที่เกี่ยวข้องกับกระบวนการจัดซื้อ มีรายการดังนี้:", size=10)

make_table(
    ["รหัสบันทึก", "ชื่อบันทึก", "ผู้จัดเก็บ", "ระยะเวลาเก็บ", "สถานที่", "ISO 9001"],
    [
        ["FR-PR-001", "ใบขอซื้อ (Purchase Requisition)", "จัดซื้อ", "3 ปี", "Server / Filing", "7.5.3"],
        ["FR-PUR-001", "ใบสั่งซื้อ (Purchase Order)", "จัดซื้อ", "3 ปี", "Server / Filing", "8.4.2"],
        ["FR-PUR-002", "แบบประเมินผู้ขาย", "จัดซื้อ", "5 ปี", "Server / Filing", "8.4.1"],
        ["FR-PUR-003", "ทะเบียนผู้ขาย (ASL)", "จัดซื้อ", "ปัจจุบัน", "Server / Filing", "8.4.1"],
        ["FR-QC-001", "รายงานตรวจสอบ QC", "QC", "5 ปี", "Server / Filing", "8.4.3"],
        ["FR-QC-002", "NCR", "QC", "5 ปี", "Server / Filing", "8.7 / 10.2"],
        ["FR-PUR-004", "Supplier Scorecard", "จัดซื้อ", "5 ปี", "Server / Filing", "8.4.1"],
        ["FR-PUR-005", "SCAR", "จัดซื้อ", "5 ปี", "Server / Filing", "10.2"],
    ],
    col_widths=[2.5, 4.5, 2.0, 2.5, 2.5, 2.5]
)

# =====================================================================
# 11. CORRECTIVE ACTIONS
# =====================================================================
doc.add_heading("11. การแก้ไขและการปรับปรุง (Corrective & Preventive Actions)", level=1)

make_table(
    ["กรณี / Event", "การดำเนินการ", "เอกสารที่เกี่ยวข้อง", "ระยะเวลา"],
    [
        ["สินค้าไม่เป็นไปตามข้อกำหนด", "แยก/ติดป้าย/คืน/ขอ Credit Note", "NCR (FR-QC-002)", "ภายใน 24 ชม."],
        ["ผู้ขายส่งของล่าช้า (>7 วัน)", "บันทึก Scorecard, ส่ง SCAR", "SCAR (FR-PUR-005)", "ภายใน 7 วัน"],
        ["ผู้ขายส่งไม่ได้คุณภาพซ้ำ", "ถอดจาก ASL ชั่วคราว, หา Supplier ใหม่", "ASL Update, SCAR", "ทันที"],
        ["ผู้ขายถูกถอดจาก ASL", "ประเมินผู้ขายใหม่แทนที่", "FR-PUR-002, FR-PUR-003", "ภายใน 30 วัน"],
        ["ข้อร้องเรียนจากลูกค้า", "ตรวจสอบย้อนกลับถึงการจัดซื้อ, แก้ไข", "Customer Complaint Log", "ตาม Wi-CR-001"],
        ["โอกาสปรับปรุง", "เสนอแนวทางลดต้นทุน/เพิ่มคุณภาพ Kaizen", "Kaizen / CI Report", "ตามรอบ CI"],
    ],
    col_widths=[4.5, 6.5, 3.5, 2.5]
)

doc.add_paragraph()

# =====================================================================
# DOCUMENT CONTROL FOOTER
# =====================================================================
doc.add_paragraph()
add_para("— จบเอกสาร (End of Document) —", bold=True, size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)

add_para(f"จัดทำโดย: ฝ่ายจัดซื้อ  |  วันที่: {date.today().strftime('%d/%m/%Y')}  |  "
         "เอกสารควบคุม: PM-001 ฉบับที่ 01",
         size=8, color=RGBColor(0x55, 0x55, 0x55),
         align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──
output = r"C:\Working_Space\OpenWork\คู่มือการสั่งซื้อ_ISO_Purchasing_Manual.docx"
doc.save(output)
print("OK")
